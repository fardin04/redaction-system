<<<<<<< HEAD
"""
🔒 Universal Redaction Tool - Streamlit Web Interface
Automated PII detection and redaction across multiple document formats
"""

import streamlit as st
import pandas as pd
from io import StringIO, BytesIO
import re

# Import redaction engine
import logic
from logic import redact_text
from accuracy import calculate_accuracy

# --------------------------
# Universal file readers
# --------------------------
def read_pdf(file):
    """Extract text from PDF files"""
    from pypdf import PdfReader
    try:
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return pd.DataFrame({"text": [text]})
    except Exception as e:
        st.error(f"❌ Error reading PDF: {e}")
        return None

def read_docx(file):
    """Extract text from DOCX files"""
    from docx import Document
    try:
        doc = Document(file)
        text = "\n".join([para.text for para in doc.paragraphs])
        return pd.DataFrame({"text": [text]})
    except Exception as e:
        st.error(f"❌ Error reading Word Doc: {e}")
        return None

def load_data(uploaded_file):
    """Universal loader: CSV, Excel, JSON, TXT, PDF, DOCX"""
    ext = uploaded_file.name.split('.')[-1].lower()

    if ext == 'csv':
        return pd.read_csv(uploaded_file)
    elif ext in ['xlsx', 'xls']:
        return pd.read_excel(uploaded_file)
    elif ext == 'json':
        try:
            return pd.read_json(uploaded_file)
        except ValueError:
            uploaded_file.seek(0)
            return pd.read_json(uploaded_file, lines=True)
    elif ext == 'txt':
        string_data = StringIO(uploaded_file.getvalue().decode("utf-8"))
        return pd.DataFrame({"text": [string_data.read()]})
    elif ext == 'pdf':
        return read_pdf(uploaded_file)
    elif ext == 'docx':
        return read_docx(uploaded_file)
    else:
        st.error(f"❌ Unsupported file format: .{ext}")
        return None

# --------------------------
# Visual redaction highlights
# --------------------------
def highlight_redaction(redacted_str, mode):
    """Highlight [ENTITY] or <ENTITY> tokens"""
    if not redacted_str:
        return ""
    
    # Highlight square brackets [NAME]
    highlighted = re.sub(
        r"(\[[A-Z0-9_]+\])",
        r"<span style='background-color:#ffe6f0; color:#b30059; font-weight:bold; padding:2px; border-radius:3px;'>\1</span>",
        redacted_str
    )
    
    # Highlight angular brackets <NAME> (for Redact mode highlights)
    highlighted = re.sub(
        r"(<[A-Z0-9_]+>)",
        r"<span style='background-color:#ffe6f0; color:#b30059; font-weight:bold; padding:2px; border-radius:3px;'>\1</span>",
        highlighted
    )
    
    return f"<pre style='white-space:pre-wrap; font-family:inherit'>{highlighted}</pre>"

# --------------------------
# Helper: Callback for Ground Truth Upload
# --------------------------
def update_gt_text():
    """Forces the text area to update when file is uploaded"""
    if st.session_state.gt_uploader is not None:
        st.session_state.acc_gt = st.session_state.gt_uploader.getvalue().decode("utf-8")

# --------------------------
# Streamlit page config
# --------------------------
st.set_page_config(
    page_title="Universal Redaction Tool",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
.stTabs [data-baseweb="tab-list"] button [data-testid="stMarkdownContainer"] p {
    font-size: 1.1em;
}
pre { background: #fafafa; padding: 10px; border-radius: 6px; }
.fixed-box { height: 200px; overflow-y: auto; padding: 8px; background: #fff; border: 1px solid #eee; border-radius:6px; }
</style>
""", unsafe_allow_html=True)

# --------------------------
# Sidebar: info only
# --------------------------
with st.sidebar:
    st.header("ℹ️ About")
    st.info("""
    Automatically detects and redacts:
    - Names & Persons
    - Email Addresses
    - Phone Numbers
    - Credit Cards, IPs, URLs
    - Dates separately
    - And more
    """)
    st.markdown("---")
    st.header("📁 Supported Formats")
    st.markdown("""
    - **Structured:** CSV, Excel, JSON
    - **Documents:** PDF, DOCX, TXT
    """)

# --------------------------
# Main tabs
# --------------------------
tab1, tab2, tab3, tab4 = st.tabs(["🚀 Live Demo", "📂 Batch Processing", "📊 Accuracy Check", "🔎 Entity Table"])

# ==========================
# TAB 1: LIVE REDACTION
# ==========================
with tab1:
    st.header("Universal PII Redaction Tool")
    st.markdown("Enter text below and watch it get redacted in real-time")

    col1, col2 = st.columns(2)

    with col1:
        st.subheader("📝 Input Text")
        txt = st.text_area(
            "Paste your text here:",
            height=350,
            value="My name is John Doe, my email is john.doe@example.com, my phone is (555) 123-4567, my ID is ID-99882 and I live in New York City, ZIP 10001. small name example: my name is alex johnson.",
            label_visibility="collapsed"
        )

        col_btn1, col_btn2 = st.columns(2)
        with col_btn1:
            if st.button("🔍 Redact Live", type="primary", use_container_width=True):
                if txt.strip():
                    # Generate versions
                    redacted_text, _, analysis, _ = redact_text(txt, mode="redact")
                    mask_text, _, _, _ = redact_text(txt, mode="mask")
                    angular_text, _, _, _ = redact_text(txt, mode="tag_angular")
                    
                    # Store results in Session State
                    st.session_state['result_redacted'] = redacted_text
                    st.session_state['result_mask'] = mask_text
                    st.session_state['result_angular'] = angular_text
                    st.session_state['original'] = txt
                    st.session_state['analysis'] = analysis 
                    st.session_state['analysis_count'] = len(analysis)

                else:
                    st.warning("⚠️ Please enter some text first!")

        with col_btn2:
            if st.button("🗑️ Clear Live", use_container_width=True):
                for k in ['result_redacted', 'result_mask', 'result_angular', 'original', 'analysis', 'analysis_count']:
                    if k in st.session_state:
                        del st.session_state[k]
                st.rerun()

    with col2:
        st.subheader("🔐 Output Controls")
        
        view_mode = st.radio("Output Mode:", options=["Redact (remove entities)", "Mask (show [ENTITY])"], index=0, horizontal=False)
        mode_value_tab1 = "redact" if view_mode.startswith("Redact") else "mask"

        st.markdown("### 🔐 Redacted Output")
        if 'result_redacted' in st.session_state:
            
            if mode_value_tab1 == "redact":
                out_text = st.session_state['result_redacted']
                vis_text = st.session_state['result_angular']
            else:
                out_text = st.session_state['result_mask']
                vis_text = st.session_state['result_mask']
            
            st.text_area("Safe Text:", value=out_text, height=200, label_visibility="collapsed")

            st.markdown("### 🎨 Visual Highlight")
            st.markdown(highlight_redaction(vis_text, mode_value_tab1), unsafe_allow_html=True)

            col_stat1, col_stat2 = st.columns(2)
            with col_stat1:
                st.metric("Items Detected", st.session_state.get('analysis_count', 0))
            with col_stat2:
                try:
                    reduction_pct = round((1 - len(st.session_state['result_redacted']) / max(1, len(st.session_state['original']))) * 100, 1)
                except Exception:
                    reduction_pct = 0.0
                st.metric("Text Reduction", f"{reduction_pct}%")

            format_choice = st.selectbox("Download format:", options=["txt", "csv", "xlsx"], index=0)
            if st.button("💾 Download Output", use_container_width=True):
                if format_choice == "txt":
                    st.download_button(label="Download TXT", data=out_text, file_name="redaction_output.txt", mime="text/plain")
                elif format_choice == "csv":
                    csv_bytes = ("text\n" + out_text.replace("\n", "\\n") + "\n").encode("utf-8")
                    st.download_button(label="Download CSV", data=csv_bytes, file_name="redaction_output.csv", mime="text/csv")
                elif format_choice == "xlsx":
                    out_df = pd.DataFrame({"text": [out_text]})
                    output = BytesIO()
                    with pd.ExcelWriter(output, engine="openpyxl") as writer:
                        out_df.to_excel(writer, index=False)
                    st.download_button(label="Download XLSX", data=output.getvalue(), file_name="redaction_output.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        else:
            st.info("👆 Enter text and click 'Redact Live' to see results")

# ==========================
# TAB 2: BATCH PROCESSING
# ==========================
with tab2:
    st.header("Batch File Processing")
    st.markdown("Upload a file and redact all PII in one go")

    col_info, col_support = st.columns(2)
    with col_info:
        st.info("Supports: **CSV, XLSX, JSON, PDF, DOCX, TXT**")

    if "batch_df" not in st.session_state:
        st.session_state.batch_df = None
    if "batch_filename" not in st.session_state:
        st.session_state.batch_filename = ""

    uploaded_file = st.file_uploader(
        "Upload your file:",
        type=['csv', 'xlsx', 'xls', 'json', 'pdf', 'docx', 'txt']
    )

    if uploaded_file:
        if uploaded_file.name != st.session_state.batch_filename:
            st.session_state.batch_df = None
            st.session_state.batch_filename = uploaded_file.name
            
        st.success(f"✅ File uploaded: **{uploaded_file.name}**")
        df_loaded = load_data(uploaded_file)
        
        if df_loaded is not None:
            st.write(f"📊 Loaded {len(df_loaded)} rows × {len(df_loaded.columns)} columns")
            
            with st.expander("👀 Preview Data", expanded=False):
                st.dataframe(df_loaded.head(3), use_container_width=True)

            options = df_loaded.columns.tolist()
            default_index = options.index("text") if "text" in options else 0
            
            target_col = st.multiselect(
                "Select column(s) to redact (pick one or more):",
                options,
                default=[options[default_index]]
            )

            if st.button("⚡ Process Entire File", type="primary", use_container_width=True):
                progress_bar = st.progress(0)
                status_text = st.empty()
                results_redacted = []
                results_mask = []
                all_batch_entities = []

                for i, row in df_loaded.iterrows():
                    combined = []
                    for c in target_col:
                        if pd.notna(row[c]):
                            combined.append(str(row[c]))
                    original_text = " ".join(combined)
                    
                    r_out, _, analysis_res, _ = redact_text(original_text, mode="redact")
                    m_out, _, _, _ = redact_text(original_text, mode="mask")
                    
                    for r in analysis_res:
                        all_batch_entities.append({
                            "Row Index": i,
                            "Entity Type": r.entity_type,
                            "Detected Text": original_text[r.start:r.end],
                            "Score": round(r.score, 2)
                        })

                    results_redacted.append(r_out)
                    results_mask.append(m_out)
                    
                    progress = (i + 1) / len(df_loaded)
                    progress_bar.progress(progress)
                    status_text.text(f"Processing: {i + 1}/{len(df_loaded)} rows...")

                status_text.text("✅ Processing complete!")
                
                df_final = df_loaded.copy()
                df_final['original_text'] = df_final[target_col].astype(str).agg(' '.join, axis=1) if len(target_col) > 1 else df_final[target_col[0]].astype(str)
                df_final['redacted_text'] = results_redacted
                df_final['mask_text'] = results_mask
                
                st.session_state.batch_df = df_final
                st.session_state['batch_entities'] = all_batch_entities

            if st.session_state.batch_df is not None:
                st.markdown("### 📋 Results Preview")
                st.dataframe(st.session_state.batch_df[['original_text', 'redacted_text', 'mask_text']].head(5), use_container_width=True)
                
                col_down_left, col_down_right = st.columns([2,3])
                with col_down_left:
                    file_format = st.selectbox("Choose download format:", options=["csv", "xlsx", "txt", "pdf"])
                
                with col_down_right:
                    data_to_download = None
                    file_name = ""
                    mime_type = ""
                    
                    if file_format == "csv":
                        data_to_download = st.session_state.batch_df.to_csv(index=False).encode('utf-8')
                        file_name = f"redacted_{st.session_state.batch_filename.split('.')[0]}.csv"
                        mime_type = "text/csv"
                    elif file_format == "xlsx":
                        output = BytesIO()
                        with pd.ExcelWriter(output, engine="openpyxl") as writer:
                            st.session_state.batch_df.to_excel(writer, index=False)
                        data_to_download = output.getvalue()
                        file_name = f"redacted_{st.session_state.batch_filename.split('.')[0]}.xlsx"
                        mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    elif file_format == "txt":
                        txt_lines = []
                        for _, r in st.session_state.batch_df.iterrows():
                            txt_lines.append(f"ORIGINAL:\n{r['original_text']}\nREDACTED:\n{r['redacted_text']}\nMASK:\n{r['mask_text']}\n\n---\n")
                        data_to_download = "\n".join(txt_lines)
                        file_name = f"redacted_{st.session_state.batch_filename.split('.')[0]}.txt"
                        mime_type = "text/plain"
                    elif file_format == "pdf":
                        try:
                            from reportlab.lib.pagesizes import letter
                            from reportlab.pdfgen import canvas
                            output = BytesIO()
                            c = canvas.Canvas(output, pagesize=letter)
                            width, height = letter
                            y = height - 40
                            for _, r in st.session_state.batch_df.iterrows():
                                lines = ["ORIGINAL: " + str(r['original_text']), "REDACTED: " + str(r['redacted_text']), "MASK: " + str(r['mask_text']), "-" * 60]
                                for line in lines:
                                    c.drawString(40, y, line[:100].replace('\n', ' '))
                                    y -= 12
                                    if y < 60:
                                        c.showPage()
                                        y = height - 40
                            c.save()
                            data_to_download = output.getvalue()
                            file_name = f"redacted_{st.session_state.batch_filename.split('.')[0]}.pdf"
                            mime_type = "application/pdf"
                        except ImportError:
                            st.error("ReportLab not installed")

                    if data_to_download:
                        st.download_button(label="💾 Download Processed File", data=data_to_download, file_name=file_name, mime=mime_type, use_container_width=True)

# ==========================
# TAB 3: ACCURACY EVALUATION
# ==========================
with tab3:
    st.header("Accuracy Evaluation")
    st.markdown("Compare redaction output against ground truth")
    st.info("Use this to evaluate redaction quality and fine-tune the system")

    col1, col2, col3 = st.columns(3)

    with col1:
        st.subheader("📌 Original Text")
        original = st.text_area("Original text:", height=200, key="acc_orig")
    
    with col2:
        st.subheader("🔐 Redacted Text")
        redacted = st.text_area("Your redaction:", height=200, key="acc_redact")
    
    with col3:
        st.subheader("🎯 Ground Truth")
        gt_file = st.file_uploader("Upload Ground Truth (.txt)", type=["txt"], key="gt_uploader", on_change=update_gt_text)
        ground_truth = st.text_area("Ground truth text:", height=130, key="acc_gt")

    if st.button("📊 Calculate Accuracy", type="primary"):
        if not original.strip() or not redacted.strip() or not ground_truth.strip():
            st.warning("⚠️ Please ensure Original, Redacted, and Ground Truth fields are filled.")
        else:
            # === FIXED: UNPACK 6 VALUES ===
            total, correct, missed, precision, recall, f1 = calculate_accuracy(original, redacted, ground_truth)

            colA, colB, colC = st.columns(3)
            colA.metric("✔️ Correctly Redacted", correct)
            colB.metric("📌 Total Entities", total)
            colC.metric("🔎 Missed Entities", missed)

            colA2, colB2, colC2 = st.columns(3)
            # Use percentage formatting
            colA2.metric("🎯 Accuracy (Recall)", f"{recall*100:.2f}%")
            colB2.metric("🎯 Precision", f"{precision*100:.2f}%")
            colC2.metric("🔥 F1 Score", f"{f1*100:.2f}%")

            st.success("🎉 Accuracy evaluation completed!")

# ==========================
# TAB 4: ENTITY TABLE
# ==========================
with tab4:
    st.header("🔎 Detected Entity Manager")
    st.markdown("View entities detected from your **Live Demo** or **Batch Processing** tasks.")

    has_live = 'analysis' in st.session_state and st.session_state['analysis']
    has_batch = 'batch_entities' in st.session_state and st.session_state['batch_entities']

    if not has_live and not has_batch:
        st.info("👋 No entities detected yet. Run a Live or Batch task first.")
    
    else:
        if has_live:
            st.subheader("⚡ Results from Live Demo")
            live_rows = []
            orig_txt = st.session_state.get('original', '')
            for r in st.session_state['analysis']:
                live_rows.append({
                    "Entity Type": r.entity_type,
                    "Detected Text": orig_txt[r.start:r.end],
                    "Score": f"{r.score:.2f}",
                    "Position": f"{r.start}-{r.end}"
                })
            df_live = pd.DataFrame(live_rows)
            st.dataframe(df_live, use_container_width=True)
            csv_live = df_live.to_csv(index=False).encode('utf-8')
            st.download_button("📥 Download Live Entities (CSV)", csv_live, "live_entities.csv", "text/csv")
            st.markdown("---")

        if has_batch:
            st.subheader("📂 Results from Batch Processing")
            df_batch = pd.DataFrame(st.session_state['batch_entities'])
            all_types = df_batch["Entity Type"].unique()
            selected_types = st.multiselect("Filter by Entity Type", options=all_types, default=all_types)
            if selected_types:
                df_filtered = df_batch[df_batch["Entity Type"].isin(selected_types)]
                st.dataframe(df_filtered, use_container_width=True)
                csv_batch = df_filtered.to_csv(index=False).encode('utf-8')
                st.download_button("📥 Download Batch Entities (CSV)", csv_batch, "batch_entities.csv", "text/csv")

# --------------------------
# Footer
# --------------------------
st.divider()
st.markdown("""
<div style='text-align: center; color: #666; padding: 20px;'>
    <p><strong>🔒 Universal Redaction Tool V1.0 <br></strong> &copy; 2025 | Developed by Super Saiyans Team</p>
</div>
=======
"""
🔒 Universal Redaction Tool - Streamlit Web Interface
Automated PII detection and redaction across multiple document formats
"""

import streamlit as st
import pandas as pd
from io import StringIO, BytesIO
import re
import datetime

# Import redaction engine
import logic
from logic import redact_text
# Ensure accuracy.py is in the same folder and has the updated 7-value return
from accuracy import calculate_accuracy

# Current year used in the footer
current_year = datetime.datetime.now().year

# --------------------------
# Universal file readers
# --------------------------
def read_pdf(file):
    """Extract text from PDF files"""
    from pypdf import PdfReader
    try:
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return pd.DataFrame({"text": [text]})
    except Exception as e:
        st.error(f"❌ Error reading PDF: {e}")
        return None

def read_docx(file):
    """Extract text from DOCX files"""
    from docx import Document
    try:
        doc = Document(file)
        text = "\n".join([para.text for para in doc.paragraphs])
        return pd.DataFrame({"text": [text]})
    except Exception as e:
        st.error(f"❌ Error reading Word Doc: {e}")
        return None

def load_data(uploaded_file):
    """Universal loader: CSV, Excel, JSON, TXT, PDF, DOCX"""
    ext = uploaded_file.name.split('.')[-1].lower()

    if ext == 'csv':
        return pd.read_csv(uploaded_file)
    elif ext in ['xlsx', 'xls']:
        return pd.read_excel(uploaded_file)
    elif ext == 'json':
        try:
            return pd.read_json(uploaded_file)
        except ValueError:
            uploaded_file.seek(0)
            return pd.read_json(uploaded_file, lines=True)
    elif ext == 'txt':
        string_data = StringIO(uploaded_file.getvalue().decode("utf-8"))
        return pd.DataFrame({"text": [string_data.read()]})
    elif ext == 'pdf':
        return read_pdf(uploaded_file)
    elif ext == 'docx':
        return read_docx(uploaded_file)
    else:
        st.error(f"❌ Unsupported file format: .{ext}")
        return None

# --------------------------
# Visual redaction highlights
# --------------------------
def highlight_redaction(redacted_str, mode):
    """Highlight [ENTITY] or <ENTITY> tokens"""
    if not redacted_str:
        return ""
    
    # Highlight square brackets [NAME]
    highlighted = re.sub(
        r"(\[[A-Z0-9_]+\])",
        r"<span style='background-color:#ffe6f0; color:#b30059; font-weight:bold; padding:2px; border-radius:3px;'>\1</span>",
        redacted_str
    )
    
    # Highlight angular brackets <NAME> (for Redact mode highlights)
    highlighted = re.sub(
        r"(<[A-Z0-9_]+>)",
        r"<span style='background-color:#ffe6f0; color:#b30059; font-weight:bold; padding:2px; border-radius:3px;'>\1</span>",
        highlighted
    )
    
    return f"<pre style='white-space:pre-wrap; font-family:inherit'>{highlighted}</pre>"

# --------------------------
# Helper: Callback for Ground Truth Upload
# --------------------------
def update_gt_text():
    """Forces the text area to update when file is uploaded"""
    if st.session_state.gt_uploader is not None:
        st.session_state.acc_gt = st.session_state.gt_uploader.getvalue().decode("utf-8")

# --------------------------
# Streamlit page config
# --------------------------
st.set_page_config(
    page_title="Universal Redaction Tool",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
.stTabs [data-baseweb="tab-list"] button [data-testid="stMarkdownContainer"] p {
    font-size: 1.1em;
}
pre { background: #fafafa; padding: 10px; border-radius: 6px; }
.fixed-box { height: 200px; overflow-y: auto; padding: 8px; background: #fff; border: 1px solid #eee; border-radius:6px; }
</style>
""", unsafe_allow_html=True)

# --------------------------
# Sidebar: info only
# --------------------------
with st.sidebar:
    st.header("ℹ️ About")
    st.info("""
    Automatically detects and redacts:
    - Names & Persons
    - Email Addresses
    - Phone Numbers
    - Credit Cards, IPs, URLs
    - Dates separately
    - And more
    """)
    st.markdown("---")
    st.header("📁 Supported Formats")
    st.markdown("""
    - **Structured:** CSV, Excel, JSON
    - **Documents:** PDF, DOCX, TXT
    """)

# --------------------------
# Main tabs
# --------------------------
tab1, tab2, tab3, tab4 = st.tabs(["🚀 Live Demo", "📂 Batch Processing", "📊 Accuracy Check", "🔎 Entity Table"])

# ==========================
# TAB 1: LIVE REDACTION
# ==========================
with tab1:
    st.header("Universal PII Redaction Tool")
    st.markdown("Enter text below and watch it get redacted in real-time")

    col1, col2 = st.columns(2)

    with col1:
        st.subheader("📝 Input Text")
        txt = st.text_area(
            "Paste your text here:",
            height=350,
            value="My name is John Doe, my email is john.doe@example.com, my phone is (555) 123-4567, my ID is ID-99882 and I live in New York City, ZIP 10001. small name example: my name is alex johnson.",
            label_visibility="collapsed"
        )

        col_btn1, col_btn2 = st.columns(2)
        with col_btn1:
            if st.button("🔍 Redact Live", type="primary", use_container_width=True):
                if txt.strip():
                    # Generate versions
                    redacted_text, _, analysis, _ = redact_text(txt, mode="redact")
                    mask_text, _, _, _ = redact_text(txt, mode="mask")
                    angular_text, _, _, _ = redact_text(txt, mode="tag_angular")
                    
                    # Store results in Session State
                    st.session_state['result_redacted'] = redacted_text
                    st.session_state['result_mask'] = mask_text
                    st.session_state['result_angular'] = angular_text
                    st.session_state['original'] = txt
                    st.session_state['analysis'] = analysis 
                    st.session_state['analysis_count'] = len(analysis)

                else:
                    st.warning("⚠️ Please enter some text first!")

        with col_btn2:
            if st.button("🗑️ Clear Live", use_container_width=True):
                for k in ['result_redacted', 'result_mask', 'result_angular', 'original', 'analysis', 'analysis_count']:
                    if k in st.session_state:
                        del st.session_state[k]
                st.rerun()

    with col2:
        st.subheader("🔐 Output Controls")
        
        view_mode = st.radio("Output Mode:", options=["Redact (remove entities)", "Mask (show [ENTITY])"], index=0, horizontal=False)
        mode_value_tab1 = "redact" if view_mode.startswith("Redact") else "mask"

        st.markdown("### 🔐 Redacted Output")
        if 'result_redacted' in st.session_state:
            
            if mode_value_tab1 == "redact":
                out_text = st.session_state['result_redacted']
                vis_text = st.session_state['result_angular']
            else:
                out_text = st.session_state['result_mask']
                vis_text = st.session_state['result_mask']
            
            st.text_area("Safe Text:", value=out_text, height=200, label_visibility="collapsed")

            st.markdown("### 🎨 Visual Highlight")
            st.markdown(highlight_redaction(vis_text, mode_value_tab1), unsafe_allow_html=True)

            col_stat1, col_stat2 = st.columns(2)
            with col_stat1:
                st.metric("Items Detected", st.session_state.get('analysis_count', 0))
            with col_stat2:
                try:
                    reduction_pct = round((1 - len(st.session_state['result_redacted']) / max(1, len(st.session_state['original']))) * 100, 1)
                except Exception:
                    reduction_pct = 0.0
                st.metric("Text Reduction", f"{reduction_pct}%")

            format_choice = st.selectbox("Download format:", options=["txt", "csv", "xlsx"], index=0)
            if st.button("💾 Download Output", use_container_width=True):
                if format_choice == "txt":
                    st.download_button(label="Download TXT", data=out_text, file_name="redaction_output.txt", mime="text/plain")
                elif format_choice == "csv":
                    csv_bytes = ("text\n" + out_text.replace("\n", "\\n") + "\n").encode("utf-8")
                    st.download_button(label="Download CSV", data=csv_bytes, file_name="redaction_output.csv", mime="text/csv")
                elif format_choice == "xlsx":
                    out_df = pd.DataFrame({"text": [out_text]})
                    output = BytesIO()
                    with pd.ExcelWriter(output, engine="openpyxl") as writer:
                        out_df.to_excel(writer, index=False)
                    st.download_button(label="Download XLSX", data=output.getvalue(), file_name="redaction_output.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        else:
            st.info("👆 Enter text and click 'Redact Live' to see results")

# ==========================
# TAB 2: BATCH PROCESSING
# ==========================
with tab2:
    st.header("Batch File Processing")
    st.markdown("Upload a file and redact all PII in one go")

    col_info, col_support = st.columns(2)
    with col_info:
        st.info("Supports: **CSV, XLSX, JSON, PDF, DOCX, TXT**")

    if "batch_df" not in st.session_state:
        st.session_state.batch_df = None
    if "batch_filename" not in st.session_state:
        st.session_state.batch_filename = ""

    uploaded_file = st.file_uploader(
        "Upload your file:",
        type=['csv', 'xlsx', 'xls', 'json', 'pdf', 'docx', 'txt']
    )

    if uploaded_file:
        if uploaded_file.name != st.session_state.batch_filename:
            st.session_state.batch_df = None
            st.session_state.batch_filename = uploaded_file.name
            
        st.success(f"✅ File uploaded: **{uploaded_file.name}**")
        df_loaded = load_data(uploaded_file)
        
        if df_loaded is not None:
            st.write(f"📊 Loaded {len(df_loaded)} rows × {len(df_loaded.columns)} columns")
            
            with st.expander("👀 Preview Data", expanded=False):
                st.dataframe(df_loaded.head(3), use_container_width=True)

            options = df_loaded.columns.tolist()
            default_index = options.index("text") if "text" in options else 0
            
            target_col = st.multiselect(
                "Select column(s) to redact (pick one or more):",
                options,
                default=[options[default_index]]
            )

            if st.button("⚡ Process Entire File", type="primary", use_container_width=True):
                progress_bar = st.progress(0)
                status_text = st.empty()
                results_redacted = []
                results_mask = []
                all_batch_entities = []

                for i, row in df_loaded.iterrows():
                    combined = []
                    for c in target_col:
                        if pd.notna(row[c]):
                            combined.append(str(row[c]))
                    original_text = " ".join(combined)
                    
                    r_out, _, analysis_res, _ = redact_text(original_text, mode="redact")
                    m_out, _, _, _ = redact_text(original_text, mode="mask")
                    
                    for r in analysis_res:
                        all_batch_entities.append({
                            "Row Index": i,
                            "Entity Type": r.entity_type,
                            "Detected Text": original_text[r.start:r.end],
                            "Score": round(r.score, 2)
                        })

                    results_redacted.append(r_out)
                    results_mask.append(m_out)
                    
                    progress = (i + 1) / len(df_loaded)
                    progress_bar.progress(progress)
                    status_text.text(f"Processing: {i + 1}/{len(df_loaded)} rows...")

                status_text.text("✅ Processing complete!")
                
                df_final = df_loaded.copy()
                df_final['original_text'] = df_final[target_col].astype(str).agg(' '.join, axis=1) if len(target_col) > 1 else df_final[target_col[0]].astype(str)
                df_final['redacted_text'] = results_redacted
                df_final['mask_text'] = results_mask
                
                st.session_state.batch_df = df_final
                st.session_state['batch_entities'] = all_batch_entities

            if st.session_state.batch_df is not None:
                st.markdown("### 📋 Results Preview")
                st.dataframe(st.session_state.batch_df[['original_text', 'redacted_text', 'mask_text']].head(5), use_container_width=True)
                
                col_down_left, col_down_right = st.columns([2,3])
                with col_down_left:
                    file_format = st.selectbox("Choose download format:", options=["csv", "xlsx", "txt", "pdf"])
                
                with col_down_right:
                    data_to_download = None
                    file_name = ""
                    mime_type = ""
                    
                    if file_format == "csv":
                        data_to_download = st.session_state.batch_df.to_csv(index=False).encode('utf-8')
                        file_name = f"redacted_{st.session_state.batch_filename.split('.')[0]}.csv"
                        mime_type = "text/csv"
                    elif file_format == "xlsx":
                        output = BytesIO()
                        with pd.ExcelWriter(output, engine="openpyxl") as writer:
                            st.session_state.batch_df.to_excel(writer, index=False)
                        data_to_download = output.getvalue()
                        file_name = f"redacted_{st.session_state.batch_filename.split('.')[0]}.xlsx"
                        mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    elif file_format == "txt":
                        txt_lines = []
                        for _, r in st.session_state.batch_df.iterrows():
                            txt_lines.append(f"ORIGINAL:\n{r['original_text']}\nREDACTED:\n{r['redacted_text']}\nMASK:\n{r['mask_text']}\n\n---\n")
                        data_to_download = "\n".join(txt_lines)
                        file_name = f"redacted_{st.session_state.batch_filename.split('.')[0]}.txt"
                        mime_type = "text/plain"
                    elif file_format == "pdf":
                        try:
                            from reportlab.lib.pagesizes import letter
                            from reportlab.pdfgen import canvas
                            output = BytesIO()
                            c = canvas.Canvas(output, pagesize=letter)
                            width, height = letter
                            y = height - 40
                            for _, r in st.session_state.batch_df.iterrows():
                                lines = ["ORIGINAL: " + str(r['original_text']), "REDACTED: " + str(r['redacted_text']), "MASK: " + str(r['mask_text']), "-" * 60]
                                for line in lines:
                                    c.drawString(40, y, line[:100].replace('\n', ' '))
                                    y -= 12
                                    if y < 60:
                                        c.showPage()
                                        y = height - 40
                            c.save()
                            data_to_download = output.getvalue()
                            file_name = f"redacted_{st.session_state.batch_filename.split('.')[0]}.pdf"
                            mime_type = "application/pdf"
                        except ImportError:
                            st.error("ReportLab not installed")

                    if data_to_download:
                        st.download_button(label="💾 Download Processed File", data=data_to_download, file_name=file_name, mime=mime_type, use_container_width=True)

# ==========================
# TAB 3: ACCURACY EVALUATION
# ==========================
with tab3:
    st.header("Accuracy Evaluation")
    st.markdown("Compare redaction output against ground truth to calculate Privacy and Utility scores.")
    st.info("Use this to evaluate both how well you hid secrets (Recall) and how readable the text remains (Utility).")

    col1, col2, col3 = st.columns(3)

    with col1:
        st.subheader("📌 Original Text")
        original = st.text_area("Original text:", height=200, key="acc_orig")
    
    with col2:
        st.subheader("🔐 Redacted Text")
        redacted = st.text_area("Your redaction:", height=200, key="acc_redact")
    
    with col3:
        st.subheader("🎯 Ground Truth")
        gt_file = st.file_uploader("Upload Ground Truth (.txt)", type=["txt"], key="gt_uploader", on_change=update_gt_text)
        ground_truth = st.text_area("Ground truth text:", height=130, key="acc_gt")

    if st.button("📊 Calculate Accuracy", type="primary"):
        if not original.strip() or not redacted.strip() or not ground_truth.strip():
            st.warning("⚠️ Please ensure Original, Redacted, and Ground Truth fields are filled.")
        else:
            # === UNPACK 7 VALUES (UPDATED for new accuracy.py) ===
            total, correct, missed, precision, recall, f1, utility = calculate_accuracy(original, redacted, ground_truth)

            # ROW 1: RAW COUNTS
            st.markdown("#### 🔢 Raw Counts")
            colA, colB, colC = st.columns(3)
            colA.metric("✔️ Correctly Redacted", correct)
            colB.metric("📌 Total Entities", total)
            colC.metric("🔎 Missed Entities", missed)

            # ROW 2: PERFORMANCE RATES
            st.markdown("#### 📈 Performance Rates")
            colA2, colB2, colC2, colD2 = st.columns(4)
            
            # Use percentage formatting with Help tooltips
            colA2.metric(
                "🎯 Recall (Privacy)", 
                f"{recall*100:.2f}%", 
                help="Privacy Score: Percentage of PII tokens successfully found and hidden."
            )
            colB2.metric(
                "🎯 Precision", 
                f"{precision*100:.2f}%", 
                help="Redaction Quality: Percentage of redacted items that were actually PII (avoiding over-redaction)."
            )
            colC2.metric(
                "⚖️ F1 Score", 
                f"{f1*100:.2f}%", 
                help="Harmonic balance between Precision and Recall."
            )
            colD2.metric(
                "📖 Utility (Similarity)", 
                f"{utility*100:.2f}%", 
                help="Structure Preservation: Measures how similar the sentence structure is to the ground truth (using Levenshtein Distance)."
            )

            st.success("🎉 Accuracy evaluation completed!")

# ==========================
# TAB 4: ENTITY TABLE
# ==========================
with tab4:
    st.header("🔎 Detected Entity Manager")
    st.markdown("View entities detected from your **Live Demo** or **Batch Processing** tasks.")

    has_live = 'analysis' in st.session_state and st.session_state['analysis']
    has_batch = 'batch_entities' in st.session_state and st.session_state['batch_entities']

    if not has_live and not has_batch:
        st.info("👋 No entities detected yet. Run a Live or Batch task first.")
    
    else:
        if has_live:
            st.subheader("⚡ Results from Live Demo")
            live_rows = []
            orig_txt = st.session_state.get('original', '')
            for r in st.session_state['analysis']:
                live_rows.append({
                    "Entity Type": r.entity_type,
                    "Detected Text": orig_txt[r.start:r.end],
                    "Score": f"{r.score:.2f}",
                    "Position": f"{r.start}-{r.end}"
                })
            df_live = pd.DataFrame(live_rows)
            st.dataframe(df_live, use_container_width=True)
            csv_live = df_live.to_csv(index=False).encode('utf-8')
            st.download_button("📥 Download Live Entities (CSV)", csv_live, "live_entities.csv", "text/csv")
            st.markdown("---")

        if has_batch:
            st.subheader("📂 Results from Batch Processing")
            df_batch = pd.DataFrame(st.session_state['batch_entities'])
            all_types = df_batch["Entity Type"].unique()
            selected_types = st.multiselect("Filter by Entity Type", options=all_types, default=all_types)
            if selected_types:
                df_filtered = df_batch[df_batch["Entity Type"].isin(selected_types)]
                st.dataframe(df_filtered, use_container_width=True)
                csv_batch = df_filtered.to_csv(index=False).encode('utf-8')
                st.download_button("📥 Download Batch Entities (CSV)", csv_batch, "batch_entities.csv", "text/csv")

# --------------------------
# Footer
# --------------------------
st.divider()
st.markdown(f"""
<div style='text-align: center; color: #666; padding: 20px;'>
    <p><strong>🔒 Universal Redaction Tool V1.0 <br></strong> &copy; {current_year} | Developed by Super Saiyans Team</p>
</div>
>>>>>>> dae57f4 (feat: Implement Universal Redaction Tool with accuracy evaluation)
""", unsafe_allow_html=True)